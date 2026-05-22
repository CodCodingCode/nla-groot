#!/usr/bin/env python3
"""Generate CoRL 2026 abstract one-pager as DOCX."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PAPER_DIR = Path(__file__).resolve().parent
FIGURE = PAPER_DIR / "figures" / "system_overview.png"
OUTPUT = PAPER_DIR / "nla_groot_abstract_corl2026.docx"

ABSTRACT = (
    "Natural-language autoencoders (NLAs) map neural activations to text and back, "
    "offering readable interfaces to model internals. We port the NLA recipe to the "
    "GR00T vision-language-action (VLA) backbone: extract per-token hidden states, "
    "train an activation verbalizer (AV) and reconstructor (AR), and inject AR outputs "
    "as live backbone steers in LIBERO simulation. Supervision uses a multimodal teacher "
    "(frames + instruction), not human descriptions of what h encodes—a confound inherited "
    "from the original LLM setting. We release nla-groot and a three-axis evaluation protocol: "
    "(1) offline reconstruction and retrieval, (2) vision-grounded caption judging, "
    "(3) closed-loop steer A/B (matched vs. mismatched language). On our main checkpoint, "
    "reconstruction metrics pass while grounding, anti-template specificity, and semantic "
    "steering fail; aggregate scores hide collapse on image_patch tokens, where retrieval "
    "margin is near zero. Claim: NL autoencoders on VLA activations are misleading without "
    "vision-grounded and behavioral checks stratified by token role—use this protocol before "
    "trusting captions or steering policies with language."
)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = notice.add_run(
        "Submitted to the 10th Conference on Robot Learning (CoRL 2026). Do not distribute."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "When Reconstruction Passes:\n"
        "Natural-Language Autoencoders on VLA Activations\n"
        "Fail Grounding and Semantic Steering"
    )
    run.bold = True
    run.font.size = Pt(14)

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run(
        "Anonymous Author(s) | Anonymous Institution | anonymous@domain.com"
    ).font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Abstract. ").bold = True
    p.add_run(ABSTRACT)

    p = doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run(
        "vision-language-action models, interpretability, activation steering, robot learning"
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Contributions.").bold = True
    for item in [
        "Port & tooling: NLA (AV/AR) on GR00T layer-16 activations with live LIBERO steering hooks.",
        "Three-axis scorecard: reconstruction/retrieval, vision-grounded judge, closed-loop steer A/B (Δ_cw).",
        "Negative result: pooled metrics pass while image_patch margin ≈ 0.003, anti-template judge 8.3%, Δ_cw = 0.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    if FIGURE.exists():
        doc.add_picture(str(FIGURE), width=Inches(3.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["Metric", "Value", "Thr.", "Verdict"]
    rows = [
        ["retrieval_margin", "0.124", "0.05", "PASS"],
        ["judge_anti_template", "0.083", "0.50", "FAIL"],
        ["sim_correct_minus_wrong", "0.000", "0.05", "WARN"],
        ["image_patch margin", "0.003", "—", "near chance"],
    ]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Takeaway. ").bold = True
    p.add_run(
        "Reconstruction certifies a text–vector codec, not pixel grounding or semantic steering. "
        "Stratify by token role before deploying NL interfaces on VLAs."
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
