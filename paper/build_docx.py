#!/usr/bin/env python3
"""Generate CoRL 2026-formatted DOCX from paper content."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PAPER_DIR = Path(__file__).resolve().parent
FIGURE = PAPER_DIR / "figures" / "system_overview.png"
OUTPUT = PAPER_DIR / "nla_groot_corl2026.docx"

REFERENCES = [
    "NVIDIA. GR00T-N1: An open foundation model for generalist humanoid robots. arXiv preprint, 2025.",
    "Connor Fraser-Taliente et al. Natural language autoencoders produce unsupervised explanations of LLM activations. Transformer Circuits Thread, 2026.",
    "Alexander Matt Turner et al. Activation addition: Steering language models without optimization. arXiv:2308.10248, 2023.",
    "Andy Zou et al. Representation engineering: A top-down approach to AI transparency. arXiv:2310.01405, 2023.",
    "Edward J. Hu et al. LoRA: Low-rank adaptation of large language models. ICLR, 2022.",
    "Anthony Brohan et al. RT-2: Vision-language-action models transfer web knowledge to robotic control. CoRL, 2023.",
    "Moo Jin Kim et al. OpenVLA: An open-source vision-language-action model. CoRL, 2024.",
]


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def add_notice(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Submitted to the 10th Conference on Robot Learning (CoRL 2026). Do not distribute."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "When Reconstruction Passes:\n"
        "Natural-Language Autoencoders on VLA Activations\n"
        "Fail Grounding and Semantic Steering"
    )
    run.bold = True
    run.font.size = Pt(14)

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run(
        "Anonymous Author(s)\n"
        "Anonymous Institution\n"
        "Anonymous City, Country\n"
        "anonymous@domain.com"
    )
    run.font.size = Pt(10)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    cap = doc.add_paragraph()
    cap_run = cap.add_run(caption)
    cap_run.bold = True
    cap_run.italic = True
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            row_cells[c_idx].text = value

    doc.add_paragraph()


def add_figure(doc: Document) -> None:
    if FIGURE.exists():
        doc.add_picture(str(FIGURE), width=Inches(6.0))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(
        "Figure 1. Pipeline. Solid: training (h from GR00T hook; gold captions from teacher). "
        "Dashed: inference (ĥ = AR(y) steers the live policy). Auditing uses AV(h); steering bypasses AV."
    )
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)


def build_docx() -> None:
    doc = Document()
    set_document_defaults(doc)
    add_notice(doc)
    add_title_block(doc)

    add_heading(doc, "Abstract", level=2)
    add_paragraph(
        doc,
        "Natural-language autoencoders (NLAs) map neural activations to text and back, offering readable "
        "interfaces to model internals. We port the NLA recipe to the GR00T vision-language-action (VLA) "
        "backbone: extract per-token hidden states, train an activation verbalizer (AV) and reconstructor "
        "(AR), and inject AR outputs as live backbone steers in LIBERO simulation. Supervision uses a "
        "multimodal teacher (frames + instruction), not human descriptions of what h encodes—a confound "
        "inherited from the original LLM setting. We release nla-groot and a three-axis evaluation protocol: "
        "(1) offline reconstruction and retrieval, (2) vision-grounded caption judging, (3) closed-loop "
        "steer A/B (matched vs. mismatched language). On our main checkpoint, reconstruction metrics pass "
        "while grounding, anti-template specificity, and semantic steering fail; aggregate scores hide collapse "
        "on image_patch tokens, where retrieval margin is near zero. Claim: NL autoencoders on VLA activations "
        "are misleading without vision-grounded and behavioral checks stratified by token role—use this "
        "protocol before trusting captions or steering policies with language.",
    )

    add_paragraph(doc, "Vision-language-action models, interpretability, activation steering, robot learning", bold_prefix="Keywords: ")

    add_heading(doc, "1 Introduction")
    add_paragraph(
        doc,
        "Vision-language-action policies such as GR00T couple a multimodal backbone to a diffusion action head. "
        "The backbone hidden state h ∈ R^2048 at layer 16 is causally upstream of actions, yet there is no "
        "standard way to read or write h in natural language. Fraser-Taliente et al. propose natural-language "
        "autoencoders: AV(h) produces a caption and AR(y) reconstructs ĥ ≈ h, yielding both an interpretability "
        "artifact and a causal handle for activation steering.",
    )
    add_paragraph(
        doc,
        "We ask whether this recipe transfers to a VLA backbone, where (i) activations are multimodal and "
        "task-conditioned, (ii) downstream behavior is the right falsification test, and (iii) training labels "
        "are synthetic: a GPT-class teacher sees camera frames and token metadata, not raw h floats. The teacher "
        "may describe scene attributes only weakly present in h; supervised fine-tuning (SFT) therefore optimizes "
        "P(teacher text | h), not faithful “what h means.”",
    )
    add_paragraph(
        doc,
        "We do not claim a strong AV or working semantic steer. We claim a negative, actionable result: "
        "autoencoder metrics can look successful while captions are vision-ungrounded and steering is "
        "behaviorally non-semantic. We release open tooling and a pre-registered three-axis scorecard so NLA "
        "ports on policies are not validated by reconstruction alone.",
        bold_prefix="Contribution. ",
    )

    add_figure(doc)

    add_heading(doc, "2 Setup (minimal)")
    add_paragraph(
        doc,
        "A forward hook on GR00T backbone_features (Qwen3-VL, layer 16, pre-action-head LayerNorm) records "
        "h_i per token at three position types: last_text, image_patch, anchor. Corpus: LIBERO Goal/Spatial/"
        "Object/10 LeRobot trajectories; α ≈ P75||h||_2 for injection scaling.",
        bold_prefix="Activations. ",
    )
    add_paragraph(
        doc,
        "Both are LoRA fine-tunes of Qwen3-4B-Instruct. AV injects α·normalize(W_p h) at a reserved slot token "
        "and is trained with CE on teacher bullets. AR reads “Summary of the following text:” and predicts "
        "ĥ/α; loss is MSE plus InfoNCE hard negatives. Joint SFT (libero_4suite_v3: 15k steps, batch 4).",
        bold_prefix="AV and AR. ",
    )
    add_paragraph(
        doc,
        "At inference, user text y maps to ĥ = AR(y) and blends into backbone tokens (image_patch or "
        "image_patch_all) on every get_action call in the official LIBERO policy server.",
        bold_prefix="Steering. ",
    )

    add_heading(doc, "3 Three-axis evaluation protocol")
    add_paragraph(
        doc,
        "Each checkpoint is scored on independent axes; overall PASS requires required metrics (pre-registered "
        "in build_v3_scorecard.py).",
    )
    add_paragraph(
        doc,
        "Closed-loop cosine cos(h, AR(AV(h))) without teacher forcing; retrieval margin (matched vs. cross-pair). "
        "Report both aggregate and per-position_type slices.",
        bold_prefix="Axis 1: Reconstruction & retrieval. ",
    )
    add_paragraph(
        doc,
        "llm_judge_av_captions.py scores captions against cached frames. We report gold (teacher) vs. AV greedy "
        "on identical rows.",
        bold_prefix="Axis 2: Vision-grounded judge. ",
    )
    add_paragraph(
        doc,
        "On put_the_bowl_on_the_plate (LIBERO Goal, GR00T checkpoint, seeds {0,1,2}): baseline, correct steer "
        "text, wrong text. Primary metric: Δ_cw = succ_correct − succ_wrong (PASS if ≥ 5pp).",
        bold_prefix="Axis 3: Closed-loop steer A/B. ",
    )

    add_heading(doc, "4 Results: metrics pass, semantics fail")
    add_table(
        doc,
        "Table 1. V3 scorecard (libero_4suite_v3). Required metrics gate overall verdict.",
        ["Metric", "Value", "Thr.", "Verdict"],
        [
            ["retrieval_margin", "0.124", "0.05", "PASS"],
            ["judge_grounding_specific", "0.417", "0.55", "WARN"],
            ["judge_anti_template", "0.083", "0.50", "FAIL"],
            ["sim_correct_minus_wrong", "0.000", "0.05", "WARN"],
        ],
    )
    add_table(
        doc,
        "Table 2. Stratified V3 metrics (libero_4suite_v3).",
        ["Position", "n", "CL cos", "Margin", "AV B%"],
        [
            ["image_patch", "85", "0.338", "0.003", "0"],
            ["last_text", "85", "0.437", "0.046", "67"],
            ["anchor", "8", "0.367†", "0.139", "75"],
            ["pooled", "178", "0.388", "0.124", "42"],
        ],
    )
    add_table(
        doc,
        "Table 3. Multimodal judge pass rates (24-row held-out sample).",
        ["Checkpoint", "Source", "Ground.", "Anti-templ."],
        [
            ["v3 (4-suite)", "gold", "0.750", "0.500"],
            ["v3 (4-suite)", "AV", "0.417", "0.083"],
        ],
    )
    add_paragraph(
        doc,
        "Baseline solves all seeds (succ = 1.0). Every steer arm fails (succ = 0) for both matching and "
        "contradictory prompts—Δ_cw = 0. Bowl displacement ≈ 0.145 m (baseline) vs. ≈ 0.07 m (all steers): "
        "intervention dampens motion uniformly rather than redirecting toward the prompted object. "
        "Steerability ≠ faithful interpretability.",
        bold_prefix="Symmetric steering (Axis 3). ",
    )

    add_heading(doc, "5 Limitations")
    add_paragraph(
        doc,
        "All experiments use a single VLA (GR00T-N1.7-LIBERO), one hook site (layer 16), and LIBERO simulation. "
        "We do not establish whether failures persist across layers, action heads, real robots, or other VLAs.",
        bold_prefix="Scope and generalization. ",
    )
    add_paragraph(
        doc,
        "Gold captions are produced by a multimodal teacher that sees pixels and metadata, not h itself. "
        "Our negative result therefore bounds teacher-supervised NLAs under this protocol.",
        bold_prefix="Teacher supervision confound. ",
    )
    add_paragraph(
        doc,
        "Axis 3 steering is tested on one LIBERO Goal task with three seeds. Judge subsamples are small for "
        "per-position_type slices.",
        bold_prefix="Evaluation coverage. ",
    )

    add_heading(doc, "6 Use this before you trust or steer")
    add_paragraph(
        doc,
        "Do not treat AV(h) as ground truth without Axis 2 on held-out frames. Report gold-vs-AV gaps per "
        "position_type; image_patch is the vision-grounding slice.",
        bold_prefix="For auditing. ",
    )
    add_paragraph(
        doc,
        "Do not deploy AR(y) backbone steers because reconstruction is high. Require Δ_cw > 0 on held-out "
        "tasks and dose calibration.",
        bold_prefix="For steering. ",
    )
    add_paragraph(
        doc,
        "Anonymous code: https://anonymous.4open.science/ (placeholder—replace before submit). "
        "Scorecard: scripts/eval/build_v3_scorecard.py.",
        bold_prefix="Reproducibility. ",
    )
    add_paragraph(
        doc,
        "NL interfaces to policy internals could help safety auditing or harm if misread as faithful; our "
        "protocol is meant to reduce overclaim.",
        bold_prefix="Broader impact. ",
    )

    add_heading(doc, "References", level=2)
    for i, ref in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph(f"[{i}] {ref}")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_docx()
